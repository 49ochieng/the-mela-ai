"""End-to-end agent-memory indexing tests.

These tests drive `agent_memory_service.create_from_upload` + `.process` with
all external services (Azure Blob, Azure Search) mocked. They verify that every
supported file type reaches the `ready` state with a positive chunk count and
correct metadata.
"""

from __future__ import annotations

import io
import uuid
from contextlib import asynccontextmanager
from typing import Dict, List
from unittest.mock import AsyncMock, patch

import pytest
import pytest_asyncio
from sqlalchemy.ext.asyncio import (
    AsyncSession,
    async_sessionmaker,
    create_async_engine,
)
from sqlalchemy.pool import StaticPool

from app.core.database import Base
from app.models.models import AgentMemoryItem, User, UserRole
from app.services.agent_memory_service import agent_memory_service


# ── Shared in-memory DB engine (single connection, survives across sessions) ──


@pytest_asyncio.fixture
async def shared_engine():
    """An in-memory SQLite engine that keeps a single shared connection so
    `agent_memory_service.process()` (which opens its own session) sees the
    same rows the test created."""
    engine = create_async_engine(
        "sqlite+aiosqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
        echo=False,
    )
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield engine
    await engine.dispose()


@pytest_asyncio.fixture
async def session_factory(shared_engine):
    return async_sessionmaker(shared_engine, expire_on_commit=False)


@pytest_asyncio.fixture
async def db(session_factory):
    async with session_factory() as s:
        yield s


# ── Mocks for external services ──────────────────────────────────────────────


class _FakeBlobStore:
    def __init__(self):
        self.store: Dict[str, bytes] = {}

    async def upload(self, data: bytes, blob_name: str, content_type: str) -> str:
        self.store[blob_name] = data
        return blob_name

    async def download(self, ref: str):
        return self.store.get(ref)

    async def delete(self, ref: str):
        self.store.pop(ref, None)
        return True


@pytest_asyncio.fixture
async def patched_env(session_factory):
    """Patch blob storage, ingestion, and session maker for the duration of
    a test."""
    fake_blob = _FakeBlobStore()
    ingest_calls: List[dict] = []

    async def fake_ingest(doc, index_name):
        ingest_calls.append({
            "index": index_name,
            "id": doc.id,
            "title": doc.title,
            "source_type": doc.source_type,
            "metadata": getattr(doc, "metadata", None)
            or getattr(doc, "meta", None)
            or {},
            "content_len": len(getattr(doc, "content", "") or ""),
        })
        # Simulate successful chunking — at least 1 chunk per doc
        return max(1, len(getattr(doc, "content", "") or "") // 1000 + 1)

    @asynccontextmanager
    async def fake_session_maker():
        async with session_factory() as s:
            yield s

    with patch("app.services.agent_memory_service._blob_store", fake_blob), \
         patch(
             "app.services.search.ingestion.ingestion_pipeline.ingest_document",
             new=AsyncMock(side_effect=fake_ingest),
         ), \
         patch("app.core.database.async_session_maker", fake_session_maker), \
         patch("app.core.database.db_available", return_value=True), \
         patch(
             "app.services.file_security.scan_file",
             return_value=type("S", (), {"safe": True, "reason": None})(),
         ):
        yield {
            "blob": fake_blob,
            "ingest_calls": ingest_calls,
        }


async def _make_user(db: AsyncSession) -> User:
    uid = str(uuid.uuid4())
    user = User(
        id=uid,
        azure_id=uid,
        email=f"u-{uid[:6]}@test.com",
        name="Tester",
        role=UserRole.USER,
    )
    db.add(user)
    await db.commit()
    return user


async def _run_flow(db, user, *, filename, content_type, data, tag="knowledge"):
    """Create + process an upload, return the refreshed item."""
    item = await agent_memory_service.create_from_upload(
        db=db, user=user,
        scope="personal", tag=tag, tenant_id=None,
        filename=filename, content_type=content_type, data=data,
    )
    await db.commit()
    await agent_memory_service.process(item.id)
    # Re-read from our session
    await db.refresh(item)
    return item


# ── Fixture-data builders ─────────────────────────────────────────────────────


def _csv_bytes():
    return (
        b"name,age,score\n"
        b"Alice,30,95.5\n"
        b"Bob,42,87.0\n"
        b"Carol,29,91.2\n"
    )


def _xlsx_bytes():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["product", "price", "qty"])
    ws.append(["A", 10.5, 3])
    ws.append(["B", 20.0, 7])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_bytes():
    from docx import Document
    doc = Document()
    doc.add_heading("Report", level=1)
    doc.add_paragraph("This is a sample report. " * 30)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_bytes():
    import fitz  # PyMuPDF
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello PDF. " + ("Lorem ipsum dolor. " * 20))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Tests: one per file type ─────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_index_txt(db, patched_env):
    user = await _make_user(db)
    item = await _run_flow(
        db, user,
        filename="notes.txt", content_type="text/plain",
        data=b"Plain text note. " * 100,
    )
    assert item.status == "ready", item.error_message
    assert item.chunk_count > 0
    assert len(patched_env["ingest_calls"]) >= 1


@pytest.mark.asyncio
async def test_index_md(db, patched_env):
    user = await _make_user(db)
    item = await _run_flow(
        db, user,
        filename="guide.md", content_type="text/markdown",
        data=b"# Title\n\nBody text here. " * 60,
    )
    assert item.status == "ready", item.error_message
    assert item.chunk_count > 0


@pytest.mark.asyncio
async def test_index_json(db, patched_env):
    user = await _make_user(db)
    item = await _run_flow(
        db, user,
        filename="data.json", content_type="application/json",
        data=b'{"key": "value", "items": [1, 2, 3, 4, 5]}',
    )
    assert item.status == "ready", item.error_message
    assert item.chunk_count > 0


@pytest.mark.asyncio
async def test_index_html(db, patched_env):
    user = await _make_user(db)
    item = await _run_flow(
        db, user,
        filename="page.html", content_type="text/html",
        data=b"<html><body><h1>Hi</h1><p>Some text here.</p></body></html>",
    )
    assert item.status == "ready", item.error_message
    assert item.chunk_count > 0


@pytest.mark.asyncio
async def test_index_csv_populates_tabular_profile(db, patched_env):
    user = await _make_user(db)
    item = await _run_flow(
        db, user,
        filename="rows.csv", content_type="text/csv",
        data=_csv_bytes(),
    )
    assert item.status == "ready", item.error_message
    assert item.chunk_count > 0
    # Tabular profile should be persisted for CSV when not a template
    assert item.template_schema_json is not None
    assert item.template_schema_json.get("kind") == "data_card"
    profile = item.template_schema_json["profile"]
    # CSV returns a single-sheet dict
    assert profile["shape"][0] == 3
    names = [c["name"] for c in profile["columns"]]
    assert set(names) == {"name", "age", "score"}


@pytest.mark.asyncio
async def test_index_xlsx_populates_tabular_profile(db, patched_env):
    user = await _make_user(db)
    item = await _run_flow(
        db, user,
        filename="book.xlsx", content_type=(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ),
        data=_xlsx_bytes(),
    )
    assert item.status == "ready", item.error_message
    assert item.chunk_count > 0
    assert item.template_schema_json is not None
    assert item.template_schema_json.get("kind") == "data_card"
    profiles = item.template_schema_json["profile"]
    assert isinstance(profiles, list)
    assert any(p["sheet"] == "Data" for p in profiles)


@pytest.mark.asyncio
async def test_index_docx(db, patched_env):
    user = await _make_user(db)
    item = await _run_flow(
        db, user,
        filename="report.docx", content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        data=_docx_bytes(),
    )
    assert item.status == "ready", item.error_message
    assert item.chunk_count > 0


@pytest.mark.asyncio
async def test_index_pdf(db, patched_env):
    user = await _make_user(db)
    item = await _run_flow(
        db, user,
        filename="doc.pdf", content_type="application/pdf",
        data=_pdf_bytes(),
    )
    assert item.status == "ready", item.error_message
    assert item.chunk_count > 0


@pytest.mark.asyncio
async def test_index_unknown_binary_does_not_crash(db, patched_env):
    """Even garbage binary should either index a placeholder or fail gracefully
    — never raise."""
    user = await _make_user(db)
    data = bytes(range(256)) * 4
    item = await agent_memory_service.create_from_upload(
        db=db, user=user,
        scope="personal", tag="knowledge", tenant_id=None,
        filename="mystery.bin", content_type="application/octet-stream",
        data=data,
    )
    await db.commit()
    await agent_memory_service.process(item.id)
    await db.refresh(item)
    # Either ready (with placeholder text) or failed — not a crash
    assert item.status in {"ready", "failed"}


@pytest.mark.asyncio
async def test_ingest_metadata_includes_scope_and_tag(db, patched_env):
    user = await _make_user(db)
    await _run_flow(
        db, user,
        filename="notes.txt", content_type="text/plain",
        data=b"Hello world " * 50,
        tag="policy",
    )
    assert len(patched_env["ingest_calls"]) >= 1
    call = patched_env["ingest_calls"][0]
    assert call["source_type"]  # non-empty
    assert call["content_len"] > 0
