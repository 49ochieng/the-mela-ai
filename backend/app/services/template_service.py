"""
Mela AI - Template Service

Parses user-uploaded DOCX or Markdown templates into a deterministic JSON
schema that the chat layer can fill in safely.

Output schema:
{
  "sections": [
    {
      "heading": str,
      "order": int,
      "placeholders": [str, ...],     # detected: {{x}}, {x}, [FILL], <<x>>
      "style_hints": str,              # bold/italic/centered/etc., free text
      "example_text": str              # the literal body of the section
    },
    ...
  ],
  "tone_summary": str,                 # heuristic — "formal", "casual", etc.
  "branding": {
    "title": str,                      # first heading or document title
    "header_text": str,                # any "header"-styled paragraph
    "footer_text": str
  }
}
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


# Regex catalogue for placeholder detection
_PLACEHOLDER_PATTERNS = [
    re.compile(r"\{\{\s*([A-Za-z0-9_\-\. ]{1,80})\s*\}\}"),  # {{name}}
    re.compile(r"\{\s*([A-Za-z0-9_\-\. ]{1,80})\s*\}"),       # {name}
    re.compile(r"\[\s*([A-Za-z0-9_\-\.\| ]{1,80})\s*\]"),     # [FILL]
    re.compile(r"<<\s*([A-Za-z0-9_\-\. ]{1,80})\s*>>"),       # <<name>>
]

# Lines starting with one of these are treated as Markdown headings.
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")

# Heuristic tone keywords for tone_summary.
_TONE_KEYWORDS = {
    "formal":   {"hereby", "pursuant", "shall", "kindly", "regards", "sincerely"},
    "casual":   {"hey", "hi", "thanks!", "cheers", "btw", "lol"},
    "marketing": {"unlock", "transform", "boost", "leverage", "discover"},
    "technical": {"api", "endpoint", "schema", "deploy", "configure"},
}


def _find_placeholders(text: str) -> List[str]:
    found: List[str] = []
    seen: set[str] = set()
    for pat in _PLACEHOLDER_PATTERNS:
        for m in pat.finditer(text):
            tag = m.group(1).strip()
            if tag and tag.lower() not in seen:
                seen.add(tag.lower())
                found.append(tag)
    return found


def _summarise_tone(text: str) -> str:
    if not text:
        return "neutral"
    lowered = text.lower()
    counts = {
        label: sum(1 for w in words if w in lowered)
        for label, words in _TONE_KEYWORDS.items()
    }
    best = max(counts.items(), key=lambda kv: kv[1])
    return best[0] if best[1] > 0 else "neutral"


# ── Markdown parsing ─────────────────────────────────────────────────────────


def _parse_markdown(text: str) -> Dict[str, Any]:
    sections: List[Dict[str, Any]] = []
    title = ""
    current_heading: Optional[str] = None
    current_buf: List[str] = []
    order = 0

    def _flush():
        nonlocal current_heading, current_buf, order
        if current_heading is None and not current_buf:
            return
        body = "\n".join(current_buf).strip()
        sections.append({
            "heading": current_heading or "Body",
            "order": order,
            "placeholders": _find_placeholders(body),
            "style_hints": "",
            "example_text": body,
        })
        order += 1
        current_heading = None
        current_buf = []

    for line in text.splitlines():
        m = _MD_HEADING.match(line.rstrip())
        if m:
            _flush()
            heading_text = m.group(2).strip()
            if not title:
                title = heading_text
            current_heading = heading_text
        else:
            current_buf.append(line)
    _flush()

    return {
        "sections": sections,
        "tone_summary": _summarise_tone(text),
        "branding": {
            "title": title,
            "header_text": "",
            "footer_text": "",
        },
    }


# ── DOCX parsing ─────────────────────────────────────────────────────────────


def _parse_docx(raw_bytes: bytes) -> Dict[str, Any]:
    try:
        from docx import Document  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx is required to parse DOCX templates. "
            "Install with: pip install python-docx"
        ) from exc

    doc = Document(io.BytesIO(raw_bytes))

    sections: List[Dict[str, Any]] = []
    title = ""
    current_heading: Optional[str] = None
    current_style: str = ""
    current_buf: List[str] = []
    order = 0

    def _flush():
        nonlocal current_heading, current_buf, current_style, order
        if current_heading is None and not current_buf:
            return
        body = "\n".join(current_buf).strip()
        sections.append({
            "heading": current_heading or "Body",
            "order": order,
            "placeholders": _find_placeholders(body),
            "style_hints": current_style,
            "example_text": body,
        })
        order += 1
        current_heading = None
        current_style = ""
        current_buf = []

    for para in doc.paragraphs:
        text = (para.text or "").rstrip()
        style_name = (para.style.name or "").lower() if para.style else ""
        if "heading" in style_name and text:
            _flush()
            current_heading = text
            current_style = style_name
            if not title:
                title = text
        else:
            current_buf.append(text)
    _flush()

    # Header / footer text (first section only)
    header_text, footer_text = "", ""
    try:
        first_section = doc.sections[0] if doc.sections else None
        if first_section:
            header_text = "\n".join(
                p.text for p in first_section.header.paragraphs if p.text
            ).strip()
            footer_text = "\n".join(
                p.text for p in first_section.footer.paragraphs if p.text
            ).strip()
    except Exception:  # pragma: no cover
        pass

    if not title and doc.core_properties.title:
        title = doc.core_properties.title

    full_text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "sections": sections,
        "tone_summary": _summarise_tone(full_text),
        "branding": {
            "title": title or "",
            "header_text": header_text,
            "footer_text": footer_text,
        },
    }


# ── Public service ───────────────────────────────────────────────────────────


class TemplateService:
    def parse(
        self,
        *,
        text: str,
        file_type: str,
        raw_bytes: Optional[bytes] = None,
    ) -> Dict[str, Any]:
        """Parse a template into the canonical schema.

        Args:
            text:       Plain-text rendering (used for MD or as a fallback).
            file_type:  'docx' | 'md' | other (other → treat as MD).
            raw_bytes:  Original bytes; required for DOCX.

        Returns:
            Schema dict (see module docstring).
        """
        ft = (file_type or "").lower()
        if ft == "docx":
            if raw_bytes is None:
                logger.warning(
                    "DOCX template missing raw_bytes — falling back to MD parser",
                )
                return _parse_markdown(text or "")
            return _parse_docx(raw_bytes)
        # Default to Markdown / text parsing
        return _parse_markdown(text or "")

    @staticmethod
    def render_prompt_block(schema: Dict[str, Any], max_chars: int = 4000) -> str:
        """Render a schema into a compact prompt block for the LLM."""
        if not schema:
            return ""
        sections = schema.get("sections") or []
        branding = schema.get("branding") or {}
        tone = schema.get("tone_summary") or "neutral"

        lines: List[str] = []
        title = branding.get("title")
        if title:
            lines.append(f"Title: {title}")
        lines.append(f"Tone: {tone}")
        lines.append("Sections (in order):")
        for s in sections:
            placeholders = s.get("placeholders") or []
            ph = f"  placeholders: {placeholders}" if placeholders else ""
            style = f"  style: {s.get('style_hints')}" if s.get("style_hints") else ""
            lines.append(f"- {s.get('order', 0)}. {s.get('heading', 'Body')}")
            if ph:
                lines.append(ph)
            if style:
                lines.append(style)
            example = (s.get("example_text") or "").strip()
            if example:
                snippet = example[:300] + ("…" if len(example) > 300 else "")
                lines.append(f"  example: {snippet}")
        if branding.get("header_text"):
            lines.append(f"Header: {branding['header_text']}")
        if branding.get("footer_text"):
            lines.append(f"Footer: {branding['footer_text']}")

        out = "\n".join(lines)
        return out[:max_chars]


template_service = TemplateService()
